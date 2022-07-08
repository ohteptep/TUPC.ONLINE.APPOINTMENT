from django.shortcuts import render, redirect
from django.http import HttpResponse 
from .models import *
from .forms import *
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from django.http import HttpResponse

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.contrib import messages

from datetime import date
from datetime import datetime

from pathlib import Path

from docx import Document
from docx2pdf import convert

import glob
import os
import shutil
import tempfile

from django.core.files import File
import pythoncom
 

# Create your views here.

name = [' ']



def sendmail_confirm(name,purp,date,depart,mail):
    print('goods')
    var = 'Good Day MR/MS '+ name + '\n'+ '\n' + 'This is to confirm your appointment for '+ purp + ' on ' + str(date) + '\n' + '\n' + 'If you have any queries feel free to reply to this email.'+'\n'+ '\n' +'Thank You!!!'+'\n'+depart+'\n'+'TUP Cavite'
    message_sent = var
    subject_sent = 'TUPC ONLINE APPOINTMENT'
    recipient_sent = [mail,]
    send_mail(subject_sent, message_sent, 'tupc.online.appointment@gmail.com', recipient_sent, fail_silently=False,)
    print('goods')

def sendmail_denied(name,purp,date,depart,mail):
    print('goods')
    var = 'Good Day MR/MS '+ name + '\n'+ '\n' +  'We would like to inform you on your appointment for' + depart + ' on ' + str(date) +  '\n'+ '\n' + 'has been declined for some reason.' + '\n'+ '\n''If you have any queries feel free to reply to this email.'+'\n' + '\n' +'Thank You!!!'+'\n'+ depart +'\n'+'TUP Cavite'
    message_sent = var
    subject_sent = 'TUPC ONLINE APPOINTMENT'
    recipient_sent = [mail,]
    send_mail(subject_sent, message_sent, 'tupc.online.appointment@gmail.com', recipient_sent, fail_silently=False,)
    print('goods')


@login_required(login_url='adminlogin')
def oaa (request):
    getalumni = AlumniBook.objects.filter(Adep = 'Office of Academic Affairs').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Office of Academic Affairs').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Office of Academic Affairs').values()

    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }
    return render (request, 'temp/oaa.html', data)

def confirm1(request,oaa_id):
    yes = AlumniBook.objects.get(id=oaa_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('oaa')
def denied1(request,oaa_id):
    yes = AlumniBook.objects.get(id=oaa_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('oaa')

def confirm2(request,oaa_id):
    yes = StudentBook.objects.get(id=oaa_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('oaa')
def denied2(request,oaa_id):
    yes = StudentBook.objects.get(id=oaa_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('oaa')

def confirm3(request,oaa_id):
    yes = GuardianBook.objects.get(id=oaa_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('oaa') 
def denied3(request,oaa_id):
    yes = GuardianBook.objects.get(id=oaa_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('oaa')

@login_required(login_url='adminlogin')
def ocr (request):
    getalumni = AlumniBook.objects.filter(Adep = 'Office of Campus Registrar').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Office of Campus Registrar').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Office of Campus Registrar').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }

    return render (request, 'temp/ocr.html', data)

def confirm4(request,ocr_id):
    yes = AlumniBook.objects.get(id=ocr_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('ocr')
def denied4(request,ocr_id):
    yes = AlumniBook.objects.get(id=ocr_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('ocr')

def confirm5(request,ocr_id):
    yes = StudentBook.objects.get(id=ocr_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('ocr')
def denied5(request,ocr_id):
    yes = StudentBook.objects.get(id=ocr_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('ocr')

def confirm6(request,ocr_id):
    yes = GuardianBook.objects.get(id=ocr_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('ocr') 
def denied6(request,ocr_id):
    yes = GuardianBook.objects.get(id=ocr_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('ocr')


def landingpage (request):
    print('landingpage')
    return render (request, 'temp/landingpage.html')

@login_required(login_url='adminlogin')
def alumni (request): 
    getalumni = AlumniBook.objects.filter(Adep = 'Alumni Office').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Alumni Office').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Alumni Office').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }

    return render (request, 'temp/alumni.html', data)
def confirm7(request,alumni_id):
    yes = AlumniBook.objects.get(id=alumni_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('alumni')
def denied7(request,alumni_id):
    yes = AlumniBook.objects.get(id=alumni_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('alumni')

def confirm8(request,alumni_id):
    yes = StudentBook.objects.get(id=alumni_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('alumni')
def denied8(request,alumni_id):
    yes = StudentBook.objects.get(id=alumni_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('alumni')

def confirm9(request,alumni_id):
    yes = GuardianBook.objects.get(id=alumni_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('alumni') 
def denied9(request,alumni_id):
    yes = GuardianBook.objects.get(id=alumni_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('alumni')

def studentlogin (request):
    sell = regacc.objects.all()
    len1 = len(sell)
    print(sell)
    print('-----------------')
    if len1 == 0:
        print('NO DATA')
    else:
        if request.method == 'POST':
            use = request.POST.get('username')
            pas = request.POST.get('pass')

            print(use)
            print(pas)
            if pas == '':
                print('change')
                return redirect ('studentlogin')
            else:
                user = authenticate(request, username=use, password=pas)
                print('userrrrrrrrrrrrrrrrr '+str(user))

                if user is not None:
                    sel1 = regacc.objects.get(username=use)
                    print('waaaaaaaaaaaaaaaa')
                    print(sel1)
                    print(sel1.Department)

                    if sel1.Department != '':
                        print('waalaaaaaa')
                        return redirect ('studentlogin')

                    login(request, user)
                    return redirect ('bookappointment')


    return render (request, 'temp/studentlogin.html')

def adminlogin (request):
    sell = regacc.objects.all()
    len1 = len(sell)
    print(sell)
    print('-----------------')
    if len1 == 0:
        print('NO DATA')
    else:
        if request.method == 'POST':
            sel = request.POST.get('choice')
            use = request.POST.get('username')
            pas = request.POST.get('passwords')

            print(sel)
            print(use)
            print(pas)
            if pas == '':
                print('change')
                return redirect ('adminlogin')
            else:
                user = authenticate(request, username=use, password=pas)
                print('userrrrrrrrrrrrrrrrr '+str(user))

                if user is not None:
                    sel1 = regacc.objects.get(username=use)
                    print('waaaaaaaaaaaaaaaa')
                    print(sel1)
                    print(sel1.Department)

                    if sel1.Department == '':
                        print('waalaaaaaa')
                        return redirect ('adminlogin')



                    login(request, user)

                    

                    if 'OAA' == sel1.Department:
                        return redirect ('oaa')
                    elif 'OCR' == sel1.Department:
                        return redirect ('ocr')
                    elif 'OGS' == sel1.Department:
                        return redirect ('ogs')
                    elif 'CL' == sel1.Department:
                        return redirect ('library')
                    elif 'OHS' == sel1.Department:
                        return redirect ('ohs')
                    elif 'UITC' == sel1.Department:
                        return redirect ('uitc')
                    elif 'USG' == sel1.Department:
                        return redirect ('usg')
                    elif 'OSA' == sel1.Department:
                        return redirect ('osa')
                    elif 'AO' == sel1.Department:
                        return redirect ('alumni')
                    elif 'SD' == sel1.Department:
                        return redirect ('securitydept')


                        
                
    return render (request, 'temp/adminlogin.html')

@login_required(login_url='studentlogin')
def bookappointment (request):
    return render (request, 'temp/bookappointment.html')

@login_required(login_url='adminlogin')
def library (request):
    getalumni = AlumniBook.objects.filter(Adep = 'Campus Library').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Campus Library').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Campus Library').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }

    return render (request, 'temp/library.html', data)

def confirm10(request,library_id):
    yes = AlumniBook.objects.get(id=library_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('library')
def denied10(request,library_id):
    yes = AlumniBook.objects.get(id=library_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('library')

def confirm11(request,library_id):
    yes = StudentBook.objects.get(id=library_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('library')
def denied11(request,library_id):
    yes = StudentBook.objects.get(id=library_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('library')

def confirm12(request,library_id):
    yes = GuardianBook.objects.get(id=library_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('library') 
def denied12(request,library_id):
    yes = GuardianBook.objects.get(id=library_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('library')

    
@login_required(login_url='adminlogin')
def ogs (request):
    getalumni = AlumniBook.objects.filter(Adep = 'Office of the Guidance Services').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Office of the Guidance Services').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Office of the Guidance Services').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }
    return render (request, 'temp/ogs.html', data)

def confirm13(request,ogs_id):
    yes = AlumniBook.objects.get(id=ogs_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('ogs')
def denied13(request,ogs_id):
    yes = AlumniBook.objects.get(id=ogs_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('ogs')

def confirm14(request,ogs_id):
    yes = StudentBook.objects.get(id=ogs_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('ogs')
def denied14(request,ogs_id):
    yes = StudentBook.objects.get(id=ogs_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('ogs')

def confirm15(request,ogs_id):
    yes = GuardianBook.objects.get(id=ogs_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('ogs') 
def denied15(request,ogs_id):
    yes = GuardianBook.objects.get(id=ogs_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('ogs')


@login_required(login_url='adminlogin')
def ohs (request):
    getalumni = AlumniBook.objects.filter(Adep = 'Office of Health Services').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Office of Health Services').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Office of Health Services').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }
    return render (request, 'temp/ohs.html', data)

def confirm16(request,ohs_id):
    yes = AlumniBook.objects.get(id=ohs_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('ohs')
def denied16(request,ohs_id):
    yes = AlumniBook.objects.get(id=ohs_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('ohs')

def confirm17(request,ohs_id):
    yes = StudentBook.objects.get(id=ohs_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('ohs')
def denied17(request,ohs_id):
    yes = StudentBook.objects.get(id=ohs_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('ohs')

def confirm18(request,ohs_id):
    yes = GuardianBook.objects.get(id=ohs_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('ohs') 
def denied18(request,ohs_id):
    yes = GuardianBook.objects.get(id=ohs_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('ohs')

@login_required(login_url='adminlogin')
def osa (request):
    getalumni = AlumniBook.objects.filter(Adep = 'Office of Student Affairs').values()
    getstudent = StudentBook.objects.filter(Sdep = 'Office of Student Affairs').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'Office of Student Affairs').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }
    return render (request, 'temp/osa.html', data)

def confirm19(request,osa_id):
    yes = AlumniBook.objects.get(id=osa_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('osa')
def denied19(request,osa_id):
    yes = AlumniBook.objects.get(id=osa_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('osa')

def confirm20(request,osa_id):
    yes = StudentBook.objects.get(id=osa_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('osa')
def denied20(request,osa_id):
    yes = StudentBook.objects.get(id=osa_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('osa')

def confirm21(request,osa_id):
    yes = GuardianBook.objects.get(id=osa_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('osa') 
def denied21(request,ohs_id):
    yes = GuardianBook.objects.get(id=osa_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('osa')

@login_required(login_url='adminlogin')
def securitydept (request):
    
    if request.method == 'POST':
        val = request.POST.get('search')
        secu1 = securitybook.objects.filter(name = val, status = 'ALUMNI' ).values()
        secu2 = securitybook.objects.filter(name = val, status = 'STUDENT' ).values()
        secu3 = securitybook.objects.filter(name = val, status = 'GUARDIAN' ).values()
    else:
        secu1 = securitybook.objects.filter(status = 'ALUMNI').values()
        secu2 = securitybook.objects.filter(status = 'STUDENT').values()
        secu3 = securitybook.objects.filter(status = 'GUARDIAN').values()

    print(secu1)
    print(secu2)
    print(secu3)

    data = {
        'alumnidata': secu1,
        'studentdata': secu2,
        'guardiandata': secu3
    }
    return render (request, 'temp/securitydept.html', data)



def signupadmin (request):
    regdata = regacc.objects.all()
    len1 = len(regdata)
    print('pinalitan')
    print(len1)
    print(regdata)
    form = adminreg()
    if len1 == 0:
        print('NO DATA')
        if request.method == 'POST':
            form = adminreg(request.POST)
            if form.is_valid():
                form.save()
                return redirect('adminlogin')
    else:
        if request.method == 'POST':

            form = adminreg(request.POST)
            if form.is_valid():
                depssss = form.cleaned_data['Department']
                print('letsssssssgoooooooo')
                print(depssss)
                for deps in regdata:
                    if deps.Department == depssss:
                        print('INVALID')
                        return redirect ('signupadmin')
                else:
                    form.save()
                    return redirect('adminlogin')

    context = {'form':form}
    return render (request, 'temp/signupadmin.html',context)

def signupstudent (request):
    regdata = regacc.objects.all()
    form = studentreg()

    if request.method == 'POST':
        form = studentreg(request.POST)
        if form.is_valid():
            form.save()
            print('na save')
            return redirect('studentlogin')
        else:
            messages.info(request, '- Username Is not Available - Password must contain: 8 Characters, 1 Capital Letter and a Digit.')
            print('eror')
            return redirect('signupstudent')


    context = {'form':form}
    return render (request, 'temp/signupstudent.html',context)


@login_required(login_url='adminlogin')
def uitc (request):
    getalumni = AlumniBook.objects.filter(Adep = 'University Information Technology Center').values()
    getstudent = StudentBook.objects.filter(Sdep = 'University Information Technology Center').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'University Information Technology Center').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }
    return render (request, 'temp/uitc.html', data)

def confirm22(request,uitc_id):
    yes = AlumniBook.objects.get(id=uitc_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('uitc')
def denied22(request,uitc_id):
    yes = AlumniBook.objects.get(id=uitc_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('uitc')

def confirm23(request,uitc_id):
    yes = StudentBook.objects.get(id=uitc_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('uitc')
def denied23(request,uitc_id):
    yes = StudentBook.objects.get(id=uitc_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('uitc')

def confirm24(request,uitc_id):
    yes = GuardianBook.objects.get(id=uitc_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('uitc') 
def denied24(request,uitc_id):
    yes = GuardianBook.objects.get(id=uitc_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('uitc')

@login_required(login_url='adminlogin')
def usg (request): 
    getalumni = AlumniBook.objects.filter(Adep = 'University Student Government').values()
    getstudent = StudentBook.objects.filter(Sdep = 'University Student Government').values()
    getguardian = GuardianBook.objects.filter(Gdep = 'University Student Government').values()


    data = {
        'alumnidata': getalumni,
        'studentdata': getstudent,
        'guardiandata': getguardian
    }
    return render (request, 'temp/usg.html', data)

def confirm25(request,usg_id):
    yes = AlumniBook.objects.get(id=usg_id)
    name = yes.Aname
    last = yes.Alast
    date = yes.Adate
    dep = yes.Adep
    purp = yes.Apurp
    status = 'ALUMNI'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('usg')
def denied25(request,usg_id):
    yes = AlumniBook.objects.get(id=usg_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Aname,yes.Apurp,yes.Adate,yes.Adep,yes.Amail)
    return redirect('usg')

def confirm26(request,usg_id):
    yes = StudentBook.objects.get(id=usg_id)
    name = yes.Sname
    last = yes.Slast
    date = yes.Sdate
    dep = yes.Sdep
    purp = yes.Spurp
    status = 'STUDENT'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('usg')
def denied26(request,usg_id):
    yes = StudentBook.objects.get(id=usg_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Sname,yes.Spurp,yes.Sdate,yes.Sdep,yes.Smail)
    return redirect('usg')

def confirm27(request,usg_id):
    yes = GuardianBook.objects.get(id=usg_id)
    name = yes.Ggname
    last = yes.Glast
    date = yes.Gdate
    dep = yes.Gdep
    purp = yes.Gpurp
    status = 'GUARDIAN'
    s = securitybook(name = name, last = last, date = date, dep = dep, purp = purp, status = status)
    s.save()
    yes.status = 'CONFIRMED'
    yes.save()
    sendmail_confirm(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('usg') 
def denied27(request,usg_id):
    yes = GuardianBook.objects.get(id=usg_id)
    yes.status = 'DENIED'
    yes.save()
    sendmail_denied(yes.Ggname,yes.Gpurp,yes.Gdate,yes.Gdep,yes.Gmail)
    return redirect('usg')

@login_required(login_url='studentlogin')
def alumniform (request):
    formAlumni = alumnus(request.POST or None)
    print('change')
    if formAlumni.is_valid():
        print('change')
        formAlumni.save()

        return redirect ('bookappointment')

    context = {
        'formAlumni': formAlumni
    }
    return render (request, 'temp/alumniform.html', context)

@login_required(login_url='studentlogin')
def studentform (request): 
    formStudent = student(request.POST or None)
    if formStudent.is_valid():
        formStudent.save()
        return redirect ('bookappointment')

    context = {
        'formStudent': formStudent
    }

    return render (request, 'temp/studentform.html', context)

@login_required(login_url='studentlogin')
def guardianform (request): 
    formGuardian = guardian(request.POST or None)
    if formGuardian.is_valid():
        formGuardian.save()
        return redirect ('bookappointment')

    context = {
        'formGuardian': formGuardian
    }

    return render (request, 'temp/guardianform.html', context)


def delete1(request,sec_id):
    yes = securitybook.objects.get(id=sec_id)
    s = securitybookFinish(name = yes.name, last = yes.last, date = yes.date, dep = yes.dep, purp = yes.purp)
    s.save()
    yes.delete()
    return redirect('securitydept')

def finish(request):
    secudata = securitybookFinish.objects.all()
    context = {
        'secufidata': secudata
    }
    return render (request, 'temp/securitydeptfinish.html', context)

def PDFile(request):
    pythoncom.CoInitialize()
    directory = os.getcwd()
    print('DIRECTORY: '+directory)
    today = date.today()
    secudata = securitybookFinish.objects.all()
    doc = Document()

    pd1 = PDFS.objects.all()
    FILENO = len(pd1) + 1

    doc.add_heading('SAVED APPOINTMENT: '+str(today), 0)
    doc.add_heading('SAVED FILE NO: '+str(FILENO), 1)
    doc.add_paragraph(" ")
    for i in range(len(secudata)):
        doc.add_paragraph("FULL NAME: "+secudata[i].name+' '+secudata[i].last)
        doc.add_paragraph("DATE OF TRANSACTION: "+secudata[i].date)
        doc.add_paragraph("DEPARTMENT: "+secudata[i].dep)
        doc.add_paragraph("PURPOSE: "+secudata[i].purp)
        doc.add_paragraph("- - - - - - - - - - - - - - - - - -")

    
    
    doc.save('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.docx')
    convert('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.docx')

    s = PDFS(PDFSave = 'SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.pdf')
    s.save()

    shutil.move('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.pdf', 'SECUPDF')
    os.remove('SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.docx')

    #os.startfile('SAVED-'+str(today)+'.pdf')
    os.startfile(directory+'\SECUPDF'+'\SAVED-FILE-'+str(FILENO)+'-'+str(today)+'.pdf')

    secudata.delete()

    return redirect('securitydept')
    


        

    
def Back(request):
    return redirect('securitydept')

def delete2(request,sec_id):
    yes = AlumniBook.objects.get(id=sec_id)
    yes.delete()
    return redirect('securitydept')

def delete3(request,sec_id):
    yes = GuardianBook.objects.get(id=sec_id)
    yes.delete()
    return redirect('securitydept')

def refresh(request):
    
    return redirect('securitydept')

def logoutuser(request):
    logout(request)
    return redirect('adminlogin')

def logoutuserwa(request):
    logout(request)
    return redirect('studentlogin')

def viewing1 (request,oaa_id):
    print('viewinggpage')
    return render (request, 'temp/viewingal.html')

def viewing2 (request,oaa_id):
    print('viewinggpage')
    return render (request, 'temp/viewingst.html')

def viewing3 (request,oaa_id):
    print('viewinggpage')
    return render (request, 'temp/viewinggd.html')

def cssform (request):
    print('viewinggpage')
    return render (request, 'temp/cssform.html')